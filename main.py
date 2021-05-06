from docx import Document
from docx.shared import Inches
from datetime import datetime

prof = input("Insira o nome do professor: ")
subject = input("Insira o nome da matéria: ")
_class = input("Insira o nome da sala: ")

day = datetime.date(datetime.now())
time = datetime.time(datetime.now())

document = Document()

document.add_heading('Lista de Presença - Turma '+_class)
p = document.add_paragraph('Essa lista foi criada no dia: '+ day.strftime("%d/%m/%Y") +' , às '+ time.strftime("%H:%M:%S"))

document.add_paragraph('Professor: ' + prof)
document.add_paragraph('Matéria: ' + subject)

document.add_heading('Lista de alunos presentes: ', level=1)
document.add_paragraph('Todos cujo os nomes estão aqui listados afirmam que estavam presentes no dia e horário descrito')

for x in range(1,36):
    document.add_paragraph(
        ' ', style='List Number'
    )

document.add_page_break()
document.save(prof + '-'+ subject+'-'+day.strftime("%d-%m-%Y") +'.docx')

